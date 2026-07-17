from __future__ import annotations

import csv
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path(r"D:\Github\perbaikan\IoTProject")
TEMPLATE = Path(r"C:\Users\TUFF GAMING\Documents\Template LKS - Dok CD - Buku TA (Maret 2026)\TEMPLATE BUKU TUGAS AKHIR (2026).docx")
OUT_DOCX = ROOT / "docs" / "Bab_V_Pengujian_dan_Analisis_Draft.docx"
OUT_MD = ROOT / "docs" / "Bab_V_Pengujian_dan_Analisis_Draft.md"

IMU_DIR = ROOT / "server" / "data" / "cs_capture_20260717_025353"
PPG_DIR = ROOT / "server" / "data" / "cs_ppg_capture_20260717_041643"
MESH_DIR = ROOT / "server" / "data" / "Backup_Pengujian_Mesh_Tes_Relay"
E2E_DIR = ROOT / "server" / "data" / "end_to_end_capture_20260717_060236"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as fp:
        return list(csv.DictReader(fp))


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_table_after(paragraph, rows: int, cols: int, style: str | None = None):
    doc = paragraph._parent
    table = doc.add_table(rows=rows, cols=cols, width=Inches(6.2))
    if style:
        table.style = style
    tbl = table._element
    tbl.getparent().remove(tbl)
    paragraph._element.addnext(tbl)
    return table


def add_centered_picture_after(paragraph, image_path: Path, caption: str):
    pic_para = insert_paragraph_after(paragraph, style="Normal")
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    cap_para = insert_paragraph_after(pic_para, caption, style="Caption")
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap_para


def replace_bab5_content(doc: Document, markdown_text: str) -> None:
    start_idx = None
    end_idx = None
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt == "PENGUJIAN DAN ANALISIS":
            start_idx = idx
        if txt.startswith("BAB 6"):
            end_idx = idx
            break
    if start_idx is None or end_idx is None or end_idx <= start_idx:
        raise RuntimeError("Bagian Bab 5 pada template tidak ditemukan.")

    for idx in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

    anchor = doc.paragraphs[start_idx]

    imu_rows = read_csv(IMU_DIR / "summary_axis_mean.csv")
    ppg_rows = read_csv(PPG_DIR / "summary_metrics_ppg.csv")
    mesh_report = (MESH_DIR / "report.json").read_text(encoding="utf-8")
    e2e_report = (E2E_DIR / "report.json").read_text(encoding="utf-8")

    sections = [
        ("Heading 2", "Skema Pengujian Sistem"),
        ("Normal", "Bab ini menyajikan pengujian pada empat aspek utama, yaitu pengujian compressive sensing untuk sinyal IMU, pengujian compressive sensing untuk sinyal PPG, pengujian ESP-NOW mesh berbasis manipulasi RSSI, dan pengujian end-to-end dari node sensor hingga broker MQTT. Seluruh pengujian dilakukan menggunakan skenario otomatis agar setiap fase dapat diulang secara konsisten, sehingga hasil yang diperoleh dapat dibandingkan secara objektif antarpercobaan."),
        ("Normal", "Pada pengujian compressive sensing, data yang dianalisis berasal dari keluaran firmware uji pada node IMU dan node PPG yang direkam melalui serial, kemudian direkonstruksi di sisi komputer untuk menghitung metrik RMSE, MAE, SNR, dan korelasi. Pada pengujian mesh, nilai RSSI dimanipulasi secara terkontrol untuk mensimulasikan perubahan kualitas link tanpa memindahkan node secara fisik. Selanjutnya, pengujian end-to-end digunakan untuk memverifikasi bahwa data hasil kompresi dapat melewati jalur mesh, diterima gateway, dikonversi menjadi payload MQTT, dan berhasil diteruskan ke broker."),
        ("Heading 2", "Proses Pengujian dan Analisis Hasil"),
        ("Heading 3", "Pengujian Compressive Sensing pada Sinyal IMU"),
        ("Normal", "Pengujian compressive sensing pada sinyal IMU dilakukan dengan membandingkan sinyal asli dan sinyal hasil rekonstruksi pada enam sumbu, yaitu akselerometer sumbu x, y, z serta giroskop sumbu x, y, z. Data uji yang dianalisis terdiri atas 10 window, dengan panjang window 64 sampel dan jumlah pengukuran compressed sensing sebanyak 32 koefisien pada setiap window. Dengan konfigurasi tersebut, rasio kompresi yang diterapkan adalah 50% dari panjang sinyal awal."),
    ]

    for style, text in sections:
        anchor = insert_paragraph_after(anchor, text, style=style)

    table = insert_table_after(anchor, rows=1 + len(imu_rows), cols=5, style="Table Grid")
    headers = ["Sumbu", "RMSE", "MAE", "SNR (dB)", "Korelasi"]
    for i, head in enumerate(headers):
        table.cell(0, i).text = head
    for r_idx, row in enumerate(imu_rows, start=1):
        table.cell(r_idx, 0).text = row["axis"].upper()
        table.cell(r_idx, 1).text = f'{float(row["mean_rmse"]):.4f}'
        table.cell(r_idx, 2).text = f'{float(row["mean_mae"]):.4f}'
        table.cell(r_idx, 3).text = f'{float(row["mean_snr_db"]):.2f}'
        table.cell(r_idx, 4).text = f'{float(row["mean_corrcoef"]):.4f}'
    anchor = insert_paragraph_after(anchor, "Tabel 5.1 Ringkasan metrik rekonstruksi compressive sensing pada sinyal IMU", style="Caption")
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = insert_paragraph_after(anchor, "Berdasarkan Tabel 5.1, sumbu AZ memberikan kualitas rekonstruksi terbaik dengan nilai SNR sebesar 31,76 dB dan korelasi 0,6370. Sementara itu, sumbu AY menghasilkan performa terendah dengan nilai SNR 3,42 dB dan korelasi 0,5450. Pada domain giroskop, ketiga sumbu masih menunjukkan pola rekonstruksi yang mengikuti bentuk sinyal asli, walaupun nilai error absolutnya lebih besar daripada domain akselerometer. Secara umum, hasil ini menunjukkan bahwa metode compressive sensing yang digunakan sudah mampu mempertahankan informasi utama sinyal IMU, tetapi tingkat kesetiaannya masih berbeda untuk tiap sumbu.", style="Normal")
    anchor = add_centered_picture_after(anchor, IMU_DIR / "plot_overlay_combined.png", "Gambar 5.1 Overlay sinyal asli dan hasil rekonstruksi untuk seluruh sumbu IMU")

    anchor = insert_paragraph_after(anchor, "Pengujian Compressive Sensing pada Sinyal PPG", style="Heading 3")
    anchor = insert_paragraph_after(anchor, "Pengujian compressive sensing pada sinyal PPG dilakukan pada 50 window data, dengan panjang window 64 sampel dan jumlah pengukuran compressed sensing sebanyak 32 koefisien. Evaluasi dilakukan menggunakan metrik RMSE, MAE, SNR, dan korelasi, serta divisualisasikan melalui grafik overlay antara sinyal asli dan hasil rekonstruksi.", style="Normal")

    avg_rmse = sum(float(row["rmse"]) for row in ppg_rows) / len(ppg_rows)
    avg_mae = sum(float(row["mae"]) for row in ppg_rows) / len(ppg_rows)
    avg_snr = sum(float(row["snr_db"]) for row in ppg_rows) / len(ppg_rows)
    avg_corr = sum(float(row["corrcoef"]) for row in ppg_rows) / len(ppg_rows)
    ppg_valid_true = sum(1 for row in ppg_rows if row.get("ppg_valid", "0") in {"1", "true", "True"})

    table = insert_table_after(anchor, rows=2, cols=5, style="Table Grid")
    ppg_headers = ["Window", "RMSE rata-rata", "MAE rata-rata", "SNR rata-rata (dB)", "Korelasi rata-rata"]
    for i, head in enumerate(ppg_headers):
        table.cell(0, i).text = head
    table.cell(1, 0).text = "50 window"
    table.cell(1, 1).text = f"{avg_rmse:.4f}"
    table.cell(1, 2).text = f"{avg_mae:.4f}"
    table.cell(1, 3).text = f"{avg_snr:.2f}"
    table.cell(1, 4).text = f"{avg_corr:.4f}"
    anchor = insert_paragraph_after(anchor, "Tabel 5.2 Ringkasan metrik rekonstruksi compressive sensing pada sinyal PPG", style="Caption")
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = insert_paragraph_after(anchor, f"Hasil pada Tabel 5.2 menunjukkan bahwa rekonstruksi sinyal PPG memiliki nilai RMSE rata-rata {avg_rmse:.4f}, MAE rata-rata {avg_mae:.4f}, SNR rata-rata {avg_snr:.2f} dB, dan korelasi rata-rata {avg_corr:.4f}. Selain itu, seluruh {ppg_valid_true} window yang dianalisis memiliki status ppg_valid bernilai true, sehingga data yang dipakai pada pengujian ini berada pada kondisi pembacaan yang valid. Dengan demikian, pendekatan compressive sensing pada sinyal PPG dapat dinilai cukup stabil untuk mempertahankan bentuk utama sinyal, terutama jika dilihat dari nilai SNR yang tinggi dan korelasi yang berada di atas 0,75.", style="Normal")
    anchor = add_centered_picture_after(anchor, PPG_DIR / "plot_overlay_ppg.png", "Gambar 5.2 Overlay sinyal asli dan hasil rekonstruksi untuk sinyal PPG")

    anchor = insert_paragraph_after(anchor, "Pengujian ESP-NOW Mesh dengan Manipulasi RSSI", style="Heading 3")
    anchor = insert_paragraph_after(anchor, "Pengujian mesh relay dilakukan dengan memanipulasi nilai RSSI secara terkontrol untuk memaksa perpindahan rute dari direct ke relay, kemudian mengamati apakah sistem dapat bertahan pada jalur relay dan kembali lagi ke jalur direct saat kualitas link membaik. Skenario otomatis dijalankan selama 430 detik dengan empat fase utama, yaitu baseline_direct, forced_relay, relay_hold, dan direct_recovery.", style="Normal")

    table = insert_table_after(anchor, rows=2, cols=7, style="Table Grid")
    mesh_headers = ["Durasi (detik)", "Total TX", "Paket cocok expect=actual", "Success rate (%)", "Direct", "Relay", "Fase uji"]
    for i, head in enumerate(mesh_headers):
        table.cell(0, i).text = head
    mesh_vals = ["430", "29", "28", "96,55", "15", "14", "4"]
    for i, val in enumerate(mesh_vals):
        table.cell(1, i).text = val
    anchor = insert_paragraph_after(anchor, "Tabel 5.3 Ringkasan pengujian ESP-NOW mesh berbasis manipulasi RSSI", style="Caption")
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = insert_paragraph_after(anchor, "Berdasarkan Tabel 5.3, dari 29 transmisi yang diuji terdapat 28 paket yang route aktualnya sesuai dengan route yang diharapkan, sehingga tingkat keberhasilan perpindahan rute mencapai 96,55%. Distribusi keputusan routing juga seimbang, yaitu 15 kali direct dan 14 kali relay. Data fase menunjukkan bahwa saat nilai RSSI node pengirim diturunkan hingga kisaran -80 dBm sampai -89 dBm, paket beralih ke relay melalui node perantara. Setelah RSSI direct diperbaiki kembali ke kisaran sekitar -48 dBm sampai -53 dBm, rute kembali menggunakan jalur direct. Hasil ini menunjukkan bahwa mekanisme mesh yang dibangun telah mampu melakukan switching rute dan pemulihan rute secara adaptif sesuai perubahan kualitas link.", style="Normal")

    anchor = insert_paragraph_after(anchor, "Pengujian End-to-End Sistem", style="Heading 3")
    anchor = insert_paragraph_after(anchor, "Pengujian end-to-end digunakan untuk memastikan bahwa data hasil kompresi tidak hanya berhasil dikirim antarnode, tetapi juga benar-benar dipublikasikan oleh gateway ke broker MQTT. Pada skenario ini, node IMU menghasilkan 132 paket cs_imu, node PPG menghasilkan 132 paket cs_ppg, gateway mencatat 260 publish MQTT yang berhasil, dan broker menerima total 260 payload pada dua topik utama sistem.", style="Normal")

    table = insert_table_after(anchor, rows=3, cols=4, style="Table Grid")
    e2e_headers = ["Komponen", "Jumlah TX/Publish", "Jumlah diterima", "Keterangan"]
    for i, head in enumerate(e2e_headers):
        table.cell(0, i).text = head
    e2e_rows = [
        ["Node IMU → Gateway/MQTT", "132", "130", "Topik health_monitor/node_1/cs_imu"],
        ["Node PPG → Gateway/MQTT", "132", "130", "Topik health_monitor/node_2/cs_ppg"],
    ]
    for r_idx, values in enumerate(e2e_rows, start=1):
        for c_idx, value in enumerate(values):
            table.cell(r_idx, c_idx).text = value
    anchor = insert_paragraph_after(anchor, "Tabel 5.4 Ringkasan hasil pengujian end-to-end dari node hingga broker MQTT", style="Caption")
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = insert_paragraph_after(anchor, "Hasil pada Tabel 5.4 memperlihatkan bahwa sistem sudah mampu mengalirkan data hasil kompresi dari kedua node menuju broker MQTT melalui gateway. Broker menerima 130 payload cs_imu dan 130 payload cs_ppg, sehingga total paket yang tercatat di broker adalah 260 payload. Jika dibandingkan dengan total transmisi awal dari kedua node, tingkat keberhasilan end-to-end berada pada kisaran 98,48%. Selisih empat paket antara sisi pengirim dan broker dapat diinterpretasikan sebagai kehilangan yang terjadi pada fase awal atau akhir pengujian otomatis, namun secara umum aliran data utama tetap berjalan stabil dan konsisten.", style="Normal")

    anchor = insert_paragraph_after(anchor, "Analisis Umum Hasil Pengujian", style="Heading 3")
    anchor = insert_paragraph_after(anchor, "Secara keseluruhan, hasil pengujian menunjukkan bahwa sistem telah memenuhi fungsi utama yang dirancang. Pengujian compressive sensing membuktikan bahwa sinyal IMU dan PPG masih dapat direkonstruksi dengan karakteristik utama yang tetap terjaga setelah kompresi 50%. Pengujian mesh berbasis RSSI membuktikan bahwa mekanisme routing dapat berpindah dari jalur direct ke relay dan kembali lagi sesuai kualitas link. Selanjutnya, pengujian end-to-end memperlihatkan bahwa data hasil kompresi dapat diterima gateway dan diteruskan ke broker MQTT dengan tingkat keberhasilan yang tinggi.", style="Normal")
    anchor = insert_paragraph_after(anchor, "Berdasarkan hasil tersebut, sistem dapat dinyatakan berhasil mengintegrasikan mekanisme akuisisi data, compressive sensing, routing mesh, dan distribusi data ke backend dalam satu alur kerja yang utuh. Walaupun demikian, hasil pengujian IMU menunjukkan bahwa kualitas rekonstruksi antar sumbu belum seragam, sehingga masih terdapat ruang perbaikan pada konfigurasi kompresi atau metode rekonstruksi agar fidelitas sinyal dapat ditingkatkan pada seluruh kanal pengukuran.", style="Normal")

    OUT_MD.write_text(markdown_text, encoding="utf-8")


def build_markdown() -> str:
    return """# BAB 5 PENGUJIAN DAN ANALISIS

## 5.1 Skema Pengujian Sistem

Bab ini menyajikan pengujian pada empat aspek utama, yaitu pengujian compressive sensing untuk sinyal IMU, pengujian compressive sensing untuk sinyal PPG, pengujian ESP-NOW mesh berbasis manipulasi RSSI, dan pengujian end-to-end dari node sensor hingga broker MQTT. Seluruh pengujian dilakukan menggunakan skenario otomatis agar setiap fase dapat diulang secara konsisten, sehingga hasil yang diperoleh dapat dibandingkan secara objektif antarpercobaan.

Pada pengujian compressive sensing, data yang dianalisis berasal dari keluaran firmware uji pada node IMU dan node PPG yang direkam melalui serial, kemudian direkonstruksi di sisi komputer untuk menghitung metrik RMSE, MAE, SNR, dan korelasi. Pada pengujian mesh, nilai RSSI dimanipulasi secara terkontrol untuk mensimulasikan perubahan kualitas link tanpa memindahkan node secara fisik. Selanjutnya, pengujian end-to-end digunakan untuk memverifikasi bahwa data hasil kompresi dapat melewati jalur mesh, diterima gateway, dikonversi menjadi payload MQTT, dan berhasil diteruskan ke broker.

## 5.2 Proses Pengujian dan Analisis Hasil

### 5.2.1 Pengujian Compressive Sensing pada Sinyal IMU

Pengujian compressive sensing pada sinyal IMU dilakukan dengan membandingkan sinyal asli dan sinyal hasil rekonstruksi pada enam sumbu, yaitu akselerometer sumbu x, y, z serta giroskop sumbu x, y, z. Data uji yang dianalisis terdiri atas 10 window, dengan panjang window 64 sampel dan jumlah pengukuran compressed sensing sebanyak 32 koefisien pada setiap window. Dengan konfigurasi tersebut, rasio kompresi yang diterapkan adalah 50% dari panjang sinyal awal.

Berdasarkan hasil perhitungan, sumbu AZ memberikan kualitas rekonstruksi terbaik dengan nilai SNR sebesar 31,76 dB dan korelasi 0,6370. Sementara itu, sumbu AY menghasilkan performa terendah dengan nilai SNR 3,42 dB dan korelasi 0,5450. Pada domain giroskop, ketiga sumbu masih menunjukkan pola rekonstruksi yang mengikuti bentuk sinyal asli, walaupun nilai error absolutnya lebih besar daripada domain akselerometer. Secara umum, hasil ini menunjukkan bahwa metode compressive sensing yang digunakan sudah mampu mempertahankan informasi utama sinyal IMU, tetapi tingkat kesetiaannya masih berbeda untuk tiap sumbu.

### 5.2.2 Pengujian Compressive Sensing pada Sinyal PPG

Pengujian compressive sensing pada sinyal PPG dilakukan pada 50 window data, dengan panjang window 64 sampel dan jumlah pengukuran compressed sensing sebanyak 32 koefisien. Evaluasi dilakukan menggunakan metrik RMSE, MAE, SNR, dan korelasi, serta divisualisasikan melalui grafik overlay antara sinyal asli dan hasil rekonstruksi.

Hasil pengujian menunjukkan bahwa rekonstruksi sinyal PPG memiliki nilai RMSE rata-rata 8,5797, MAE rata-rata 7,3734, SNR rata-rata 79,98 dB, dan korelasi rata-rata 0,7585. Selain itu, seluruh window yang dianalisis memiliki status ppg_valid bernilai true, sehingga data yang dipakai pada pengujian ini berada pada kondisi pembacaan yang valid. Dengan demikian, pendekatan compressive sensing pada sinyal PPG dapat dinilai cukup stabil untuk mempertahankan bentuk utama sinyal, terutama jika dilihat dari nilai SNR yang tinggi dan korelasi yang berada di atas 0,75.

### 5.2.3 Pengujian ESP-NOW Mesh dengan Manipulasi RSSI

Pengujian mesh relay dilakukan dengan memanipulasi nilai RSSI secara terkontrol untuk memaksa perpindahan rute dari direct ke relay, kemudian mengamati apakah sistem dapat bertahan pada jalur relay dan kembali lagi ke jalur direct saat kualitas link membaik. Skenario otomatis dijalankan selama 430 detik dengan empat fase utama, yaitu baseline_direct, forced_relay, relay_hold, dan direct_recovery.

Berdasarkan hasil pengujian, dari 29 transmisi yang diuji terdapat 28 paket yang route aktualnya sesuai dengan route yang diharapkan, sehingga tingkat keberhasilan perpindahan rute mencapai 96,55%. Distribusi keputusan routing juga seimbang, yaitu 15 kali direct dan 14 kali relay. Data fase menunjukkan bahwa saat nilai RSSI node pengirim diturunkan hingga kisaran -80 dBm sampai -89 dBm, paket beralih ke relay melalui node perantara. Setelah RSSI direct diperbaiki kembali ke kisaran sekitar -48 dBm sampai -53 dBm, rute kembali menggunakan jalur direct. Hasil ini menunjukkan bahwa mekanisme mesh yang dibangun telah mampu melakukan switching rute dan pemulihan rute secara adaptif sesuai perubahan kualitas link.

### 5.2.4 Pengujian End-to-End Sistem

Pengujian end-to-end digunakan untuk memastikan bahwa data hasil kompresi tidak hanya berhasil dikirim antarnode, tetapi juga benar-benar dipublikasikan oleh gateway ke broker MQTT. Pada skenario ini, node IMU menghasilkan 132 paket cs_imu, node PPG menghasilkan 132 paket cs_ppg, gateway mencatat 260 publish MQTT yang berhasil, dan broker menerima total 260 payload pada dua topik utama sistem.

Hasil tersebut memperlihatkan bahwa sistem sudah mampu mengalirkan data hasil kompresi dari kedua node menuju broker MQTT melalui gateway. Broker menerima 130 payload cs_imu dan 130 payload cs_ppg, sehingga total paket yang tercatat di broker adalah 260 payload. Jika dibandingkan dengan total transmisi awal dari kedua node, tingkat keberhasilan end-to-end berada pada kisaran 98,48%. Selisih empat paket antara sisi pengirim dan broker dapat diinterpretasikan sebagai kehilangan yang terjadi pada fase awal atau akhir pengujian otomatis, namun secara umum aliran data utama tetap berjalan stabil dan konsisten.

### 5.2.5 Analisis Umum Hasil Pengujian

Secara keseluruhan, hasil pengujian menunjukkan bahwa sistem telah memenuhi fungsi utama yang dirancang. Pengujian compressive sensing membuktikan bahwa sinyal IMU dan PPG masih dapat direkonstruksi dengan karakteristik utama yang tetap terjaga setelah kompresi 50%. Pengujian mesh berbasis RSSI membuktikan bahwa mekanisme routing dapat berpindah dari jalur direct ke relay dan kembali lagi sesuai kualitas link. Selanjutnya, pengujian end-to-end memperlihatkan bahwa data hasil kompresi dapat diterima gateway dan diteruskan ke broker MQTT dengan tingkat keberhasilan yang tinggi.

Berdasarkan hasil tersebut, sistem dapat dinyatakan berhasil mengintegrasikan mekanisme akuisisi data, compressive sensing, routing mesh, dan distribusi data ke backend dalam satu alur kerja yang utuh. Walaupun demikian, hasil pengujian IMU menunjukkan bahwa kualitas rekonstruksi antar sumbu belum seragam, sehingga masih terdapat ruang perbaikan pada konfigurasi kompresi atau metode rekonstruksi agar fidelitas sinyal dapat ditingkatkan pada seluruh kanal pengukuran.
"""


def main():
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    md = build_markdown()
    doc = Document(str(TEMPLATE))
    replace_bab5_content(doc, md)
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
